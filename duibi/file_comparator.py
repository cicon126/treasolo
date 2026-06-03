import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import difflib
import os
import threading
from datetime import datetime


class FileComparatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("文件对比工具")
        self.root.geometry("1200x800")
        
        self.file1_path = tk.StringVar()
        self.file2_path = tk.StringVar()
        self.file1_content = ""
        self.file2_content = ""
        self.diff_result = []
        self.is_loading = False
        self.browse_btn1 = None
        self.browse_btn2 = None
        self.compare_btn = None
        
        self.setup_ui()

    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        file_frame = ttk.LabelFrame(main_frame, text="文件选择", padding="10")
        file_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        file_frame.columnconfigure(1, weight=1)
        file_frame.columnconfigure(3, weight=1)
        
        ttk.Label(file_frame, text="文件 1:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(file_frame, textvariable=self.file1_path, width=80).grid(row=0, column=1, padx=5, pady=5, sticky=(tk.W, tk.E))
        self.browse_btn1 = ttk.Button(file_frame, text="浏览...", command=self.browse_file1)
        self.browse_btn1.grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(file_frame, text="文件 2:").grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        ttk.Entry(file_frame, textvariable=self.file2_path, width=80).grid(row=1, column=1, padx=5, pady=5, sticky=(tk.W, tk.E))
        self.browse_btn2 = ttk.Button(file_frame, text="浏览...", command=self.browse_file2)
        self.browse_btn2.grid(row=1, column=2, padx=5, pady=5)
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, columnspan=2, pady=(0, 10))
        
        self.compare_btn = ttk.Button(button_frame, text="对比文件", command=self.compare_files)
        self.compare_btn.grid(row=0, column=0, padx=10)
        ttk.Button(button_frame, text="导出结果", command=self.export_result).grid(row=0, column=1, padx=10)
        ttk.Button(button_frame, text="清空内容", command=self.clear_all).grid(row=0, column=2, padx=10)
        
        content_frame = ttk.Frame(main_frame)
        content_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        content_frame.columnconfigure(0, weight=1)
        content_frame.columnconfigure(1, weight=1)
        content_frame.rowconfigure(1, weight=1)
        
        file1_label = ttk.Label(content_frame, text="文件 1 内容")
        file1_label.grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        
        file2_label = ttk.Label(content_frame, text="文件 2 内容")
        file2_label.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        self.text1 = ScrolledText(content_frame, wrap=tk.NONE, width=60, height=20)
        self.text1.grid(row=1, column=0, padx=5, pady=5, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.text1.tag_configure("delete", background="#ffcccc", foreground="#ff0000")
        self.text1.tag_configure("empty", background="#f0f0f0")
        
        self.text2 = ScrolledText(content_frame, wrap=tk.NONE, width=60, height=20)
        self.text2.grid(row=1, column=1, padx=5, pady=5, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.text2.tag_configure("insert", background="#ccffcc", foreground="#009900")
        self.text2.tag_configure("empty", background="#f0f0f0")
        
        result_frame = ttk.LabelFrame(main_frame, text="对比结果统计", padding="10")
        result_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))
        result_frame.columnconfigure(0, weight=1)
        
        self.result_text = ScrolledText(result_frame, wrap=tk.WORD, height=8)
        self.result_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.result_text.tag_configure("header", font=("Arial", 10, "bold"))
        self.result_text.tag_configure("info", foreground="#0066cc")
        self.result_text.tag_configure("add", foreground="#009900")
        self.result_text.tag_configure("delete", foreground="#ff0000")
        self.result_text.tag_configure("change", foreground="#ff9900")

        self.status_var = tk.StringVar(value="请选择要对比的两个文件")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0))

    FILETYPES = [
        ("所有支持的文件", "*.txt *.py *.csv *.json *.xml *.doc *.docx"),
        ("Word文档", "*.doc *.docx"),
        ("文本文件", "*.txt"),
        ("Python文件", "*.py"),
        ("CSV文件", "*.csv"),
        ("JSON文件", "*.json"),
        ("XML文件", "*.xml"),
        ("所有文件", "*.*")
    ]

    def set_ui_state(self, enabled):
        state = 'normal' if enabled else 'disabled'
        if self.browse_btn1:
            self.browse_btn1.config(state=state)
        if self.browse_btn2:
            self.browse_btn2.config(state=state)
        if self.compare_btn:
            self.compare_btn.config(state=state)
        self.root.update_idletasks()

    def browse_file1(self):
        if self.is_loading:
            return
        filename = filedialog.askopenfilename(
            title="选择文件 1",
            filetypes=self.FILETYPES
        )
        if filename:
            self.file1_path.set(filename)
            self.load_file_async(filename, 1)

    def browse_file2(self):
        if self.is_loading:
            return
        filename = filedialog.askopenfilename(
            title="选择文件 2",
            filetypes=self.FILETYPES
        )
        if filename:
            self.file2_path.set(filename)
            self.load_file_async(filename, 2)

    def load_file_async(self, filepath, file_num):
        if self.is_loading:
            return
        self.is_loading = True
        self.set_ui_state(False)
        self.status_var.set(f"正在加载文件 {file_num}，请稍候...")
        
        thread = threading.Thread(
            target=self._load_file_thread,
            args=(filepath, file_num)
        )
        thread.daemon = True
        thread.start()

    def _load_file_thread(self, filepath, file_num):
        ext = os.path.splitext(filepath)[1].lower()
        content = ""
        error_msg = None

        try:
            if ext in ('.docx', '.doc'):
                content, error_msg = self._read_word_file(filepath, ext)
            else:
                content, error_msg = self._read_text_file(filepath)
        except Exception as e:
            error_msg = str(e)
            content = None

        self.root.after(0, lambda: self._on_file_loaded(content, file_num, filepath, error_msg))

    def _on_file_loaded(self, content, file_num, filepath, error_msg):
        self.is_loading = False
        self.set_ui_state(True)

        if error_msg:
            messagebox.showerror("错误", f"读取文件失败: {error_msg}")
            self.status_var.set("加载失败")
            return

        if content is None:
            self.status_var.set("加载失败")
            return

        if file_num == 1:
            self.file1_content = content
            self.text1.delete(1.0, tk.END)
            self.text1.insert(1.0, content)
        else:
            self.file2_content = content
            self.text2.delete(1.0, tk.END)
            self.text2.insert(1.0, content)

        self.status_var.set(f"文件 {file_num} 加载成功: {filepath}")

    def _read_text_file(self, filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return (f.read(), None)
        except UnicodeDecodeError:
            try:
                with open(filepath, 'r', encoding='gbk') as f:
                    return (f.read(), None)
            except Exception as e:
                return (None, f"无法读取文件 {filepath}: {str(e)}")
        except Exception as e:
            return (None, f"无法读取文件 {filepath}: {str(e)}")

    def _read_word_file(self, filepath, ext):
        if ext == '.docx':
            return self._read_docx(filepath)
        else:
            return self._read_doc(filepath)

    def _read_docx(self, filepath):
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text
                if text.strip() == '':
                    paragraphs.append('')
                else:
                    paragraphs.append(text)

            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    tables_text.append(' | '.join(row_data))

            all_text = paragraphs
            if tables_text:
                all_text.append("")
                all_text.append("--- 表格内容 ---")
                all_text.extend(tables_text)

            return ('\n'.join(all_text), None)
        except ImportError:
            return (None, "读取 .docx 文件需要 python-docx 库。\n请运行: pip install python-docx")
        except Exception as e:
            return (None, f"无法读取 .docx 文件 {filepath}: {str(e)}")

    def _read_doc(self, filepath):
        content, error = self._read_doc_via_com(filepath)
        if content is not None:
            return (content, None)
        
        alt_content, alt_error = self._read_doc_via_olefile(filepath)
        if alt_content is not None:
            self.root.after(0, lambda: messagebox.showwarning(
                "提示",
                f"Word COM 方式读取失败，已使用备用方式读取。\n\n"
                f"COM 错误: {error[:100]}...\n\n"
                f"建议：将 .doc 文件另存为 .docx 格式以获得更好的兼容性。"
            ))
            return (alt_content, None)
        
        combined_error = (
            f"无法读取 .doc 文件，两种方式均失败：\n\n"
            f"方式1 (Word COM): {error}\n\n"
            f"方式2 (直接解析): {alt_error}\n\n"
            f"建议：\n"
            f"1. 在 Word 中打开文件并另存为 .docx 格式\n"
            f"2. 确认文件未损坏且不是重命名的其他格式"
        )
        return (None, combined_error)

    def _read_doc_via_com(self, filepath):
        word = None
        try:
            import win32com.client
            import pythoncom
            import time

            pythoncom.CoInitialize()
            
            try:
                word = win32com.client.GetActiveObject("Word.Application")
            except Exception:
                try:
                    word = win32com.client.Dispatch("Word.Application")
                except Exception as e:
                    return (None, self._get_word_error_msg(str(e), filepath))
            
            word.Visible = False
            word.DisplayAlerts = 0
            
            max_retries = 2
            retry_delay = 0.5
            doc = None
            
            for attempt in range(max_retries):
                try:
                    abs_path = os.path.abspath(filepath)
                    doc = word.Documents.Open(
                        abs_path,
                        ReadOnly=True,
                        AddToRecentFiles=False,
                        Visible=False
                    )
                    break
                except Exception as e:
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                    else:
                        self._safe_quit_word(word)
                        word = None
                        return (None, self._get_word_error_msg(str(e), filepath))
            
            if doc is None:
                self._safe_quit_word(word)
                return (None, f"无法打开 .doc 文件: {filepath}")
            
            try:
                content = doc.Content.Text
                if content:
                    content = content.replace('\r', '\n')
            finally:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
            
            self._safe_quit_word(word)
            word = None
            
            return (content, None)
            
        except ImportError as e:
            return (None, f"缺少 pywin32 库: {str(e)}")
        except Exception as e:
            self._safe_quit_word(word)
            return (None, self._get_word_error_msg(str(e), filepath))
        finally:
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass

    def _read_doc_via_olefile(self, filepath):
        try:
            import olefile
            
            if not olefile.isOleFile(filepath):
                return (None, "文件不是有效的 OLE 复合文档格式，可能已损坏或不是真正的 .doc 文件")
            
            ole = olefile.OleFileIO(filepath)
            
            text_parts = []
            
            if ole.exists('WordDocument'):
                try:
                    doc_stream = ole.openstream('WordDocument')
                    doc_data = doc_stream.read()
                    
                    try:
                        text = doc_data.decode('utf-16-le', errors='ignore')
                    except:
                        try:
                            text = doc_data.decode('gbk', errors='ignore')
                        except:
                            text = doc_data.decode('latin-1', errors='ignore')
                    
                    clean_text = []
                    for line in text.split('\n'):
                        line = line.replace('\x00', '').strip()
                        if line:
                            clean_text.append(line)
                    
                    if clean_text:
                        text_parts.extend(clean_text)
                except Exception as e:
                    pass
            
            if ole.exists('1Table') or ole.exists('0Table'):
                try:
                    table_stream = ole.openstream('1Table' if ole.exists('1Table') else '0Table')
                    table_data = table_stream.read()
                    try:
                        table_text = table_data.decode('utf-16-le', errors='ignore')
                    except:
                        table_text = table_data.decode('gbk', errors='ignore')
                    
                    table_clean = []
                    for line in table_text.split('\n'):
                        line = line.replace('\x00', '').strip()
                        if line and len(line) > 3 and not line.startswith(''):
                            table_clean.append(line)
                    
                    if table_clean and len(table_clean) < 50:
                        if text_parts:
                            text_parts.append("")
                            text_parts.append("--- 表格/附加内容 ---")
                        text_parts.extend(table_clean)
                except:
                    pass
            
            ole.close()
            
            if not text_parts:
                return (None, "无法从文件中提取文本内容，文件可能加密或格式特殊。\n建议在 Word 中另存为 .docx 格式后使用。")
            
            result = '\n'.join(text_parts)
            
            if len(result.strip()) < 10:
                return (None, "提取的内容过少，可能文件格式特殊。\n建议在 Word 中另存为 .docx 格式后使用。")
            
            return (result, None)
            
        except ImportError:
            return (None, "缺少 olefile 库。\n请运行: pip install olefile")
        except Exception as e:
            return (None, f"解析 .doc 文件失败: {str(e)}\n建议在 Word 中另存为 .docx 格式后使用。")

    def _safe_quit_word(self, word):
        if word is None:
            return
        try:
            for doc in word.Documents:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
            word.Quit()
        except:
            pass

    def _get_word_error_msg(self, error_msg, filepath):
        error_lower = error_msg.lower()
        
        if "server execution failed" in error_lower or "服务器运行失败" in error_msg or "-2146959355" in error_msg:
            return (
                f"Word COM 服务器启动失败 (错误码 -2146959355)。\n\n"
                f"可能的解决方案：\n"
                f"1. 确认已安装 Microsoft Word 并能正常启动\n"
                f"2. 打开 Word，接受任何首次启动提示后再试\n"
                f"3. 在任务管理器中结束所有 WINWORD.EXE 进程后重试\n"
                f"4. 运行 Office 安装修复程序\n"
                f"5. 强烈建议：将 .doc 文件另存为 .docx 格式后使用（无需依赖 Word）\n\n"
                f"文件: {filepath}\n"
                f"详细错误: {error_msg}"
            )
        elif "无法创建" in error_msg or "create" in error_lower or "dispatch" in error_lower:
            return (
                f"无法启动 Word 应用程序。\n\n"
                f"可能的解决方案：\n"
                f"1. 确认已安装 Microsoft Word\n"
                f"2. Word 版本与 Python 位数匹配（32位/64位）\n"
                f"3. 强烈建议：将 .doc 文件另存为 .docx 格式后使用\n\n"
                f"详细错误: {error_msg}"
            )
        else:
            return (
                f"读取 .doc 文件时出错。\n\n"
                f"建议：将 .doc 文件在 Word 中另存为 .docx 格式后使用，\n"
                f".docx 格式兼容性更好且无需依赖 Word COM。\n\n"
                f"文件: {filepath}\n"
                f"详细错误: {error_msg}"
            )

    def compare_files(self):
        if not self.file1_path.get() or not self.file2_path.get():
            messagebox.showwarning("提示", "请先选择两个文件！")
            return
        
        if not self.file1_content or not self.file2_content:
            messagebox.showwarning("提示", "文件内容为空，请重新选择！")
            return
        
        self.status_var.set("正在对比文件...")
        self.root.update()
        
        lines1 = self.file1_content.splitlines(keepends=True)
        lines2 = self.file2_content.splitlines(keepends=True)
        
        diff = list(difflib.ndiff(lines1, lines2))
        
        self.diff_result = diff
        self.display_diff(diff, lines1, lines2)
        self.display_summary(diff)
        
        self.status_var.set("文件对比完成！")

    def display_diff(self, diff, lines1, lines2):
        self.text1.delete(1.0, tk.END)
        self.text2.delete(1.0, tk.END)
        
        line_num1 = 1
        line_num2 = 1
        
        i = 0
        while i < len(diff):
            line = diff[i]
            code = line[0]
            content = line[2:] if len(line) > 2 else ""
            
            if code == ' ':
                self.text1.insert(tk.END, content)
                self.text2.insert(tk.END, content)
                line_num1 += 1
                line_num2 += 1
            elif code == '-':
                if i + 1 < len(diff) and diff[i+1][0] == '?':
                    self.text1.insert(tk.END, content, "delete")
                    self.text2.insert(tk.END, self._make_placeholder(content), "empty")
                    line_num1 += 1
                elif i + 1 < len(diff) and diff[i+1][0] == '+':
                    if i + 2 < len(diff) and diff[i+2][0] == '?':
                        self.text1.insert(tk.END, content, "delete")
                        self.text2.insert(tk.END, diff[i+1][2:], "insert")
                        i += 2
                        line_num1 += 1
                        line_num2 += 1
                    else:
                        self.text1.insert(tk.END, content, "delete")
                        self.text2.insert(tk.END, diff[i+1][2:], "insert")
                        i += 1
                        line_num1 += 1
                        line_num2 += 1
                else:
                    self.text1.insert(tk.END, content, "delete")
                    self.text2.insert(tk.END, self._make_placeholder(content), "empty")
                    line_num1 += 1
            elif code == '+':
                self.text1.insert(tk.END, self._make_placeholder(content), "empty")
                self.text2.insert(tk.END, content, "insert")
                line_num2 += 1
            i += 1

    def _make_placeholder(self, content):
        if content.endswith('\n'):
            return '\n'
        return ''

    def display_summary(self, diff):
        self.result_text.delete(1.0, tk.END)
        
        added = 0
        deleted = 0
        changed = 0
        
        i = 0
        while i < len(diff):
            line = diff[i]
            code = line[0]
            
            if code == '?':
                i += 1
                continue
            
            if code == '-':
                has_following_question = i + 1 < len(diff) and diff[i+1][0] == '?'
                has_following_plus = i + 1 < len(diff) and diff[i+1][0] == '+'
                has_following_question_after_plus = i + 2 < len(diff) and diff[i+2][0] == '?'
                
                if has_following_question:
                    changed += 1
                    i += 2
                    continue
                elif has_following_plus:
                    changed += 1
                    if has_following_question_after_plus:
                        i += 3
                    else:
                        i += 2
                    continue
                else:
                    deleted += 1
            elif code == '+':
                has_following_question = i + 1 < len(diff) and diff[i+1][0] == '?'
                if has_following_question:
                    i += 2
                    continue
                else:
                    added += 1
            i += 1
        
        total_lines1 = len(self.file1_content.splitlines())
        total_lines2 = len(self.file2_content.splitlines())
        
        self.result_text.insert(tk.END, "=== 文件对比结果 ===\n", "header")
        self.result_text.insert(tk.END, f"对比时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n", "info")
        self.result_text.insert(tk.END, f"文件 1: {self.file1_path.get()}\n", "info")
        self.result_text.insert(tk.END, f"  总行数: {total_lines1}\n", "info")
        self.result_text.insert(tk.END, f"文件 2: {self.file2_path.get()}\n", "info")
        self.result_text.insert(tk.END, f"  总行数: {total_lines2}\n\n", "info")
        self.result_text.insert(tk.END, "--- 差异统计 ---\n", "header")
        self.result_text.insert(tk.END, f"新增行数: ", "header")
        self.result_text.insert(tk.END, f"{added}\n", "add")
        self.result_text.insert(tk.END, f"删除行数: ", "header")
        self.result_text.insert(tk.END, f"{deleted}\n", "delete")
        self.result_text.insert(tk.END, f"修改行数: ", "header")
        self.result_text.insert(tk.END, f"{changed}\n\n", "change")
        self.result_text.insert(tk.END, f"总差异数: {added + deleted + changed}\n\n", "header")
        
        if added + deleted + changed == 0:
            self.result_text.insert(tk.END, "✓ 两个文件内容完全相同！\n", "add")
        else:
            self.result_text.insert(tk.END, "--- 详细差异 ---\n\n", "header")
            self._display_detailed_diff(diff)

    def _display_detailed_diff(self, diff):
        line_num1 = 0
        line_num2 = 0
        diff_group = []
        
        i = 0
        while i < len(diff):
            line = diff[i]
            code = line[0]
            content = line[2:] if len(line) > 2 else ""
            
            if code == ' ':
                if diff_group:
                    self._write_diff_group(diff_group, line_num1, line_num2)
                    diff_group = []
                line_num1 += 1
                line_num2 += 1
            elif code == '-':
                line_num1 += 1
                if diff_group and diff_group[-1][0] == '+':
                    self._write_diff_group(diff_group, line_num1 - 1, line_num2)
                    diff_group = [(code, content, line_num1, line_num2)]
                else:
                    diff_group.append((code, content, line_num1, line_num2))
            elif code == '+':
                line_num2 += 1
                diff_group.append((code, content, line_num1, line_num2))
            i += 1
        
        if diff_group:
            self._write_diff_group(diff_group, line_num1, line_num2)

    def _write_diff_group(self, group, ln1, ln2):
        has_delete = any(item[0] == '-' for item in group)
        has_add = any(item[0] == '+' for item in group)
        
        if has_delete and has_add:
            self.result_text.insert(tk.END, f"[修改] 行 {group[0][2] if group[0][0] == '-' else group[0][3]}\n", "change")
        elif has_delete:
            self.result_text.insert(tk.END, f"[删除] 行 {group[0][2]}\n", "delete")
        else:
            self.result_text.insert(tk.END, f"[新增] 行 {group[0][3]}\n", "add")
        
        for code, content, ln1_item, ln2_item in group:
            if code == '-':
                self.result_text.insert(tk.END, f"  - {content.rstrip()}\n", "delete")
            elif code == '+':
                self.result_text.insert(tk.END, f"  + {content.rstrip()}\n", "add")
        
        self.result_text.insert(tk.END, "\n")

    def export_result(self):
        if not self.diff_result:
            messagebox.showwarning("提示", "请先进行文件对比！")
            return
        
        default_name = f"对比结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filename = filedialog.asksaveasfilename(
            title="导出对比结果",
            defaultextension=".txt",
            initialfile=default_name,
            filetypes=[("文本文件", "*.txt"), ("HTML文件", "*.html")]
        )
        
        if filename:
            try:
                if filename.endswith('.html'):
                    self._export_html(filename)
                else:
                    self._export_text(filename)
                messagebox.showinfo("成功", f"对比结果已导出到:\n{filename}")
                self.status_var.set(f"导出成功: {filename}")
            except Exception as e:
                messagebox.showerror("错误", f"导出失败: {str(e)}")

    def _export_text(self, filename):
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("=" * 60 + "\n")
            f.write("文件对比结果\n")
            f.write("=" * 60 + "\n")
            f.write(f"对比时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"文件 1: {self.file1_path.get()}\n")
            f.write(f"文件 2: {self.file2_path.get()}\n")
            f.write("=" * 60 + "\n\n")
            
            for line in self.diff_result:
                f.write(line)

    def _export_html(self, filename):
        added = 0
        deleted = 0
        changed = 0
        i = 0
        while i < len(self.diff_result):
            line = self.diff_result[i]
            code = line[0] if len(line) > 0 else ' '
            
            if code == '?':
                i += 1
                continue
            
            if code == '-':
                has_following_question = i + 1 < len(self.diff_result) and len(self.diff_result[i+1]) > 0 and self.diff_result[i+1][0] == '?'
                has_following_plus = i + 1 < len(self.diff_result) and len(self.diff_result[i+1]) > 0 and self.diff_result[i+1][0] == '+'
                has_following_question_after_plus = i + 2 < len(self.diff_result) and len(self.diff_result[i+2]) > 0 and self.diff_result[i+2][0] == '?'
                
                if has_following_question:
                    changed += 1
                    i += 2
                    continue
                elif has_following_plus:
                    changed += 1
                    if has_following_question_after_plus:
                        i += 3
                    else:
                        i += 2
                    continue
                else:
                    deleted += 1
            elif code == '+':
                has_following_question = i + 1 < len(self.diff_result) and len(self.diff_result[i+1]) > 0 and self.diff_result[i+1][0] == '?'
                if has_following_question:
                    i += 2
                    continue
                else:
                    added += 1
            i += 1
        
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>文件对比结果</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #333; }}
        .info {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
        .stats {{ display: flex; gap: 20px; margin: 20px 0; }}
        .stat-box {{ padding: 15px; border-radius: 5px; text-align: center; }}
        .add {{ background: #e8f5e9; color: #2e7d32; }}
        .delete {{ background: #ffebee; color: #c62828; }}
        .change {{ background: #fff3e0; color: #ef6c00; }}
        pre {{ background: #fafafa; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        .line {{ display: block; }}
        .line-insert {{ background: #ccffcc; }}
        .line-delete {{ background: #ffcccc; }}
        .line-equal {{ background: #ffffff; }}
    </style>
</head>
<body>
    <h1>文件对比结果</h1>
    <div class="info">
        <p><strong>对比时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p><strong>文件 1:</strong> {self.file1_path.get()}</p>
        <p><strong>文件 2:</strong> {self.file2_path.get()}</p>
    </div>
    <div class="stats">
        <div class="stat-box add"><strong>新增</strong><br>{added} 行</div>
        <div class="stat-box delete"><strong>删除</strong><br>{deleted} 行</div>
        <div class="stat-box change"><strong>修改</strong><br>{changed} 行</div>
    </div>
    <h2>详细对比</h2>
    <pre>
"""
        
        for line in self.diff_result:
            escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if line.startswith('+ '):
                html_content += f'<span class="line line-insert">{escaped_line}</span>'
            elif line.startswith('- '):
                html_content += f'<span class="line line-delete">{escaped_line}</span>'
            else:
                html_content += f'<span class="line line-equal">{escaped_line}</span>'
        
        html_content += """
    </pre>
</body>
</html>
"""
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def clear_all(self):
        self.file1_path.set("")
        self.file2_path.set("")
        self.file1_content = ""
        self.file2_content = ""
        self.diff_result = []
        self.text1.delete(1.0, tk.END)
        self.text2.delete(1.0, tk.END)
        self.result_text.delete(1.0, tk.END)
        self.status_var.set("请选择要对比的两个文件")


def main():
    root = tk.Tk()
    app = FileComparatorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
